# -*- coding: utf-8 -*-
"""
Lectura - A self-hosted Markdown note-taking application

26. Type Hints - Comprehensive type annotations
27. Error Handling Standardization - Custom exception classes
28. Configuration Validation - Pydantic models for config
29. Logging - Structured logging throughout
"""
import io
import json
import logging
import os
import shutil
import urllib.request
def _get_httpx():
    import httpx
    return httpx
from pathlib import Path
from typing import Dict, List, Optional, Any

from fastapi import FastAPI, HTTPException, UploadFile, File, Request
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, HttpUrl, Field, field_validator

# ── 29. Logging Setup ──────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)
logger.info("Starting Lectura")

# ── lazy-loaded optional heavy deps ───────────────────────────────────────────
# These are imported on first use to speed up startup.

def _get_fitz():
    try:
        import fitz
        return fitz
    except ImportError:
        return None

def _get_docx():
    try:
        from docx import Document as DocxDocument
        return DocxDocument
    except ImportError:
        return None

def _get_git():
    try:
        import git
        return git
    except ImportError:
        return None

def _get_weasyprint():
    try:
        from weasyprint import HTML as WP_HTML
        return WP_HTML
    except ImportError:
        return None


def _get_gdrive():
    try:
        from google.oauth2.credentials import Credentials
        from google_auth_oauthlib.flow import Flow
        from googleapiclient.discovery import build as gdrive_build
        from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
        return {
            "Credentials": Credentials,
            "Flow": Flow,
            "build": gdrive_build,
            "MediaIoBaseUpload": MediaIoBaseUpload,
            "MediaIoBaseDownload": MediaIoBaseDownload,
        }
    except ImportError:
        return None

# ── PDF to Markdown conversion ────────────────────────────────────────────────

import re as _re

def _pdf_extract_image(page, img_info, fitz, images_dir: Path, prefix: str, counter: int) -> Optional[str]:
    """Extract a single image from a PDF page and save it. Returns markdown image ref or None."""
    try:
        xref = img_info[0]
        base_image = fitz.Pixmap(page.parent, xref)
        # Convert CMYK to RGB
        if base_image.n > 4:
            base_image = fitz.Pixmap(fitz.csRGB, base_image)
        ext = "png"
        img_bytes = base_image.tobytes(ext)
        fname = f"{prefix}_img_{counter}.{ext}"
        dest = images_dir / fname
        dest.write_bytes(img_bytes)
        return f"![image](/static/images/{fname})"
    except Exception:
        return None


def _pdf_table_to_markdown(table) -> str:
    """Convert a PyMuPDF table object to a markdown table string."""
    rows = table.extract()
    if not rows or len(rows) < 1:
        return ""
    # Clean cells: replace None with empty string, strip whitespace and newlines
    clean_rows = []
    for row in rows:
        clean_rows.append([str(cell).replace("\n", " ").strip() if cell else "" for cell in row])
    if not clean_rows:
        return ""
    # Build markdown table
    lines = []
    header = clean_rows[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in clean_rows[1:]:
        # Pad row to match header length
        while len(row) < len(header):
            row.append("")
        lines.append("| " + " | ".join(row[:len(header)]) + " |")
    return "\n".join(lines)


def _pdf_to_markdown(fitz, data: bytes) -> str:
    """Convert PDF bytes to structured Markdown with text, images, and tables."""
    doc = fitz.open(stream=data, filetype="pdf")
    images_dir = BASE / "static" / "images"
    images_dir.mkdir(exist_ok=True)

    # Check if table extraction is available (PyMuPDF >= 1.23.0)
    has_tables = False
    try:
        if len(doc) > 0:
            doc[0].find_tables()
            has_tables = True
    except (AttributeError, Exception):
        pass

    # First pass: collect all font sizes to determine heading thresholds
    font_sizes: list[float] = []
    for page in doc:
        blocks = page.get_text("dict", flags=getattr(fitz, 'TEXT_PRESERVE_WHITESPACE', 1))["blocks"]
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    size = round(span["size"], 1)
                    text = span["text"].strip()
                    if text:
                        font_sizes.append(size)

    if not font_sizes and not doc.page_count:
        return ""

    # Determine body size (most common) and heading thresholds
    from collections import Counter
    body_size = 0.0
    heading_map: dict[float, int] = {}
    if font_sizes:
        size_counts = Counter(font_sizes)
        body_size = size_counts.most_common(1)[0][0]
        unique_sizes = sorted(set(s for s in font_sizes if s > body_size + 1.5), reverse=True)
        for i, size in enumerate(unique_sizes[:3]):
            heading_map[size] = i + 1

    # Generate a unique prefix for images from this import
    import time
    img_prefix = f"pdf_{int(time.time())}"
    img_counter = 0

    # Second pass: build markdown
    md_lines: list[str] = []
    for page_idx, page in enumerate(doc):
        if page_idx > 0:
            md_lines.append("\n---\n")

        # Extract tables for this page (to know which regions to skip in text)
        table_rects: list = []
        page_tables: list[str] = []
        if has_tables:
            try:
                tables = page.find_tables()
                for table in tables.tables:
                    table_rects.append(fitz.Rect(table.bbox))
                    md_table = _pdf_table_to_markdown(table)
                    if md_table:
                        page_tables.append((table.bbox[1], md_table))  # (y-position, markdown)
            except Exception:
                pass

        # Extract images for this page
        page_images: list[tuple] = []
        try:
            for img_info in page.get_images(full=True):
                img_counter += 1
                md_img = _pdf_extract_image(page, img_info, fitz, images_dir, img_prefix, img_counter)
                if md_img:
                    # Get image position on page via xref lookup
                    xref = img_info[0]
                    try:
                        img_rects = page.get_image_rects(xref)
                        y_pos = img_rects[0].y0 if img_rects else 9999
                    except Exception:
                        y_pos = 9999
                    page_images.append((y_pos, md_img))
        except Exception:
            pass

        # Collect text blocks with y-position, skipping table regions
        text_elements: list[tuple] = []
        blocks = page.get_text("dict", flags=getattr(fitz, 'TEXT_PRESERVE_WHITESPACE', 1))["blocks"]
        for block in blocks:
            if block.get("type") != 0:
                continue

            # Skip text that falls inside a table region
            block_rect = fitz.Rect(block["bbox"])
            in_table = False
            for tr in table_rects:
                if block_rect.intersects(tr):
                    in_table = True
                    break
            if in_table:
                continue

            block_text_parts: list[str] = []
            block_heading_level = 0

            for line in block.get("lines", []):
                line_parts: list[str] = []
                for span in line.get("spans", []):
                    text = span["text"]
                    # Fix common encoding artifacts
                    text = text.replace("\ufffd", "'")  # replacement char → apostrophe
                    text = text.replace("\u00a0", " ")  # non-breaking space → space
                    text = text.replace("\xad", "")      # soft hyphen
                    if not text.strip():
                        if text:
                            line_parts.append(" ")
                        continue

                    size = round(span["size"], 1)
                    flags = span.get("flags", 0)
                    is_bold = bool(flags & 2**4)
                    is_italic = bool(flags & 2**1)

                    if size in heading_map:
                        block_heading_level = max(block_heading_level, heading_map[size])

                    stripped = text.strip()
                    leading = text[:len(text) - len(text.lstrip())]
                    trailing = text[len(text.rstrip()):]

                    if is_bold and is_italic:
                        chunk = f"{leading}***{stripped}***{trailing}"
                    elif is_bold:
                        chunk = f"{leading}**{stripped}**{trailing}"
                    elif is_italic:
                        chunk = f"{leading}*{stripped}*{trailing}"
                    else:
                        chunk = text

                    line_parts.append(chunk)

                line_text = "".join(line_parts).rstrip()
                # Merge adjacent bold markers: **word1** **word2** → **word1 word2**
                line_text = _re.sub(r'\*\*\s*\*\*', ' ', line_text)
                line_text = _re.sub(r'\*\*\*\s*\*\*\*', ' ', line_text)
                # Clean up empty bold/italic markers
                line_text = line_text.replace("****", "").replace("******", "")
                if line_text:
                    block_text_parts.append(line_text)

            # Join lines within a block — detect whether lines are a flowing
            # paragraph (should be joined with spaces) or intentionally separate
            # lines (lists, short lines, etc.)
            if len(block_text_parts) <= 1:
                paragraph = "".join(block_text_parts).strip()
            else:
                # Heuristic: if most lines end without sentence-ending
                # punctuation and are long-ish, they're a wrapped paragraph.
                block_width = block["bbox"][2] - block["bbox"][0]
                page_width = page.rect.width
                # Lines that reach close to the block width are likely wrapped
                avg_len = sum(len(l) for l in block_text_parts) / len(block_text_parts)
                short_lines = sum(1 for l in block_text_parts[:-1]  # skip last line
                                 if len(l) < avg_len * 0.5)
                is_list = all(_re.match(r'^[\u2022\u2023\u25E6\u2043\-\*]\s|^\d+[\.\)]\s', l)
                              for l in block_text_parts)
                # If it looks like a list or most lines are short, preserve breaks
                if is_list or short_lines > len(block_text_parts) * 0.4:
                    paragraph = "\n".join(block_text_parts).strip()
                else:
                    paragraph = " ".join(block_text_parts).strip()
            if not paragraph:
                continue

            # Clean up any remaining encoding issues
            paragraph = _re.sub(r'\s{2,}', ' ', paragraph) if "\n" not in paragraph else paragraph

            y_pos = block["bbox"][1]
            if _re.match(r'^[\u2022\u2023\u25E6\u2043\-\*]\s', paragraph):
                paragraph = _re.sub(r'^[\u2022\u2023\u25E6\u2043]', '-', paragraph)
                text_elements.append((y_pos, paragraph))
            elif _re.match(r'^\d+[\.\)]\s', paragraph):
                text_elements.append((y_pos, paragraph))
            elif block_heading_level:
                prefix = "#" * block_heading_level
                clean = paragraph.replace("**", "").replace("***", "")
                text_elements.append((y_pos, f"{prefix} {clean}"))
            else:
                text_elements.append((y_pos, paragraph))

        # Merge text, images, and tables sorted by vertical position on the page
        all_elements = text_elements + page_images + page_tables
        all_elements.sort(key=lambda x: x[0])
        for _, content in all_elements:
            md_lines.append(content)

    result = "\n\n".join(md_lines)
    result = _re.sub(r'\n{4,}', '\n\n\n', result)
    return result.strip()


# ── app setup ──────────────────────────────────────────────────────────────────
BASE = Path(__file__).parent
CONFIG_PATH = BASE / "config.json"

# Workspace directory — defaults to ./notes, can be changed at runtime
def _init_notes_dir() -> Path:
    """Load last workspace from config, or default to BASE/notes."""
    cfg_path = BASE / "config.json"
    if cfg_path.exists():
        try:
            cfg = json.loads(cfg_path.read_text())
            workspace = cfg.get("workspace")
            if workspace:
                p = Path(workspace)
                if p.is_dir():
                    return p
        except Exception:
            pass
    default = BASE / "notes"
    default.mkdir(exist_ok=True)
    return default

NOTES = _init_notes_dir()

# GitHub OAuth
GITHUB_CLIENT_ID = os.getenv('GITHUB_CLIENT_ID', 'Ov23li70jsJUucF7xlgH')
GITHUB_CLIENT_SECRET = os.getenv('GITHUB_CLIENT_SECRET', '5abf8eeaacb82592b5498063858815a1cb3e20ba')
GITHUB_TOKEN_PATH = BASE / ".github_token.json"


# Google Drive OAuth
GDRIVE_CLIENT_ID = os.getenv("GDRIVE_CLIENT_ID", "337821946016-jml660ds3tvf7ecdpevulg8s56aupe3h.apps.googleusercontent.com")
GDRIVE_CLIENT_SECRET = os.getenv("GDRIVE_CLIENT_SECRET", "")
GDRIVE_TOKEN_PATH = BASE / ".gdrive_token.json"
GDRIVE_SCOPES = ["https://www.googleapis.com/auth/drive.file"]
GDRIVE_CLIENT_CONFIG = {
    "web": {
        "client_id": GDRIVE_CLIENT_ID,
        "client_secret": GDRIVE_CLIENT_SECRET,
        "auth_uri": "https://accounts.google.com/o/oauth2/auth",
        "token_uri": "https://oauth2.googleapis.com/token",
        "redirect_uris": ["http://localhost:8000/gdrive/callback"],
    }
}

app = FastAPI(title="Lectura", docs_url=None, redoc_url=None)
app.mount("/static", StaticFiles(directory=BASE / "static"), name="static")


def load_config() -> dict:
    if CONFIG_PATH.exists():
        return json.loads(CONFIG_PATH.read_text())
    return {}


def save_config(cfg: dict):
    CONFIG_PATH.write_text(json.dumps(cfg, indent=2))


# ── 27. Error Handling Standardization ────────────────────────────────────────
class AppError(HTTPException):
    """Base exception for application errors."""
    def __init__(self, status_code: int, detail: str, error_code: str = None):
        self.error_code = error_code or f"ERR_{status_code}"
        super().__init__(status_code=status_code, detail=detail)


class NotFoundError(AppError):
    """Resource not found."""
    def __init__(self, resource: str, identifier: str):
        super().__init__(404, f"{resource} '{identifier}' not found", "NOT_FOUND")


class ValidationError(AppError):
    """Validation error."""
    def __init__(self, detail: str):
        super().__init__(400, detail, "VALIDATION_ERROR")


class ConfigurationError(AppError):
    """Configuration error."""
    def __init__(self, detail: str):
        super().__init__(400, detail, "CONFIG_ERROR")


class DependencyError(AppError):
    """Missing dependency error."""
    def __init__(self, dependency: str, install_hint: str = ""):
        detail = f"{dependency} not installed"
        if install_hint:
            detail += f" — run: {install_hint}"
        super().__init__(501, detail, "DEPENDENCY_ERROR")


# ── 28. Configuration Validation with Pydantic ────────────────────────────────
class GitHubConfig(BaseModel):
    """GitHub configuration model."""
    repo_url: Optional[str] = ""
    branch: str = "main"
    token: Optional[str] = ""
    
    @field_validator('branch')
    @classmethod
    def validate_branch(cls, v: str) -> str:
        if v and not v.replace('-', '').replace('_', '').isalnum():
            raise ValueError('Invalid branch name')
        return v or "main"



class GDriveConfig(BaseModel):
    """Google Drive configuration model."""
    enabled: bool = False
    connected: bool = False


class AppConfig(BaseModel):
    """Application configuration model."""
    github: GitHubConfig = Field(default_factory=GitHubConfig)
gdrive: GDriveConfig = Field(default_factory=GDriveConfig)


# ── models ─────────────────────────────────────────────────────────────────────
class NoteBody(BaseModel):
    """Request body for saving notes."""
    content: str = Field(..., min_length=0, max_length=10_000_000)  # 10MB max


class PublishBody(BaseModel):
    """Request body for publishing notes."""
    content: str
    message: str = Field(default="", max_length=1000)


class HtmlBody(BaseModel):
    """Request body for HTML export."""
    html: str


class ConfigBody(BaseModel):
    """Request body for configuration updates."""
    config: Dict[str, Any]


class FileListResponse(BaseModel):
    """Response model for file listing."""
    files: List[str]
    folders: List[str]
    current_folder: str = ""


class SearchResult(BaseModel):
    """Single search result."""
    name: str
    snippet: str


class SearchResponse(BaseModel):
    """Response model for search."""
    results: List[SearchResult]


def validate_path(name: str) -> Path:
    """Resolve a user-supplied path and ensure it stays within NOTES directory."""
    resolved = (NOTES / name).resolve()
    if not str(resolved).startswith(str(NOTES.resolve())):
        raise HTTPException(403, "Path traversal not allowed")
    return resolved


# ── routes ─────────────────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def root():
    return FileResponse(BASE / "static" / "index.html")


# ── workspace management ───────────────────────────────────────────────────────

class WorkspaceBody(BaseModel):
    path: str

@app.get("/workspace")
async def get_workspace():
    return {"path": str(NOTES.resolve()), "name": NOTES.resolve().name}

@app.post("/workspace")
async def set_workspace(body: WorkspaceBody):
    global NOTES
    folder = Path(body.path).expanduser().resolve()
    if not folder.exists():
        raise HTTPException(400, f"Folder does not exist: {body.path}")
    if not folder.is_dir():
        raise HTTPException(400, f"Not a directory: {body.path}")
    NOTES = folder
    # Persist in config
    cfg = load_config()
    cfg["workspace"] = str(folder)
    save_config(cfg)
    logger.info(f"Workspace changed to: {folder}")
    return {"path": str(folder), "name": folder.name}

@app.post("/workspace/browse")
async def browse_directory(body: WorkspaceBody = None):
    """List subdirectories of a given path for folder browsing."""
    target = Path(body.path).expanduser().resolve() if body and body.path else Path.home()
    if not target.exists() or not target.is_dir():
        target = Path.home()
    dirs = []
    try:
        for item in sorted(target.iterdir()):
            if item.is_dir() and not item.name.startswith("."):
                dirs.append({"name": item.name, "path": str(item)})
    except PermissionError:
        pass
    return {"current": str(target), "parent": str(target.parent), "dirs": dirs}


# ── file management ────────────────────────────────────────────────────────────

@app.get("/files")
async def list_files(folder: str = "", recursive: bool = True):
    """List files and folders. If recursive=True, returns all files/folders for tree view.
    If recursive=False, returns only items in the specified folder."""
    # Determine the target directory
    if folder:
        target_dir = NOTES / folder
        if not target_dir.exists() or not target_dir.is_dir():
            raise HTTPException(404, f"Folder '{folder}' not found")
    else:
        target_dir = NOTES
    
    if recursive:
        # Return ALL files and folders recursively for tree view
        files = []
        folders = []
        
        for p in NOTES.rglob("*"):
            if p.is_file() and p.suffix == ".md":
                rel_path = p.relative_to(NOTES).as_posix()
                # Get first line preview
                preview = ""
                try:
                    with open(p, 'r', encoding='utf-8') as f:
                        first_line = f.readline().strip()
                        # Remove markdown formatting for preview
                        preview = first_line.lstrip('#').lstrip('*').lstrip('-').lstrip('>').strip()
                        if len(preview) > 60:
                            preview = preview[:60] + "..."
                except:
                    preview = ""
                files.append({"path": rel_path, "preview": preview})
            elif p.is_dir():
                folders.append(p.relative_to(NOTES).as_posix() + "/")
        
        files = sorted(files, key=lambda x: x["path"])
        folders = sorted(folders)
        return {"files": files, "folders": folders, "current_folder": ""}
    else:
        # Non-recursive: only items in the specified folder
        files = []
        for p in target_dir.glob("*.md"):
            if p.is_file():
                rel_path = p.relative_to(NOTES).as_posix()
                preview = ""
                try:
                    with open(p, 'r', encoding='utf-8') as f:
                        first_line = f.readline().strip()
                        preview = first_line.lstrip('#').lstrip('*').lstrip('-').lstrip('>').strip()
                        if len(preview) > 60:
                            preview = preview[:60] + "..."
                except:
                    preview = ""
                files.append({"path": rel_path, "preview": preview})
        files = sorted(files, key=lambda x: x["path"])
        
        folders = []
        for d in target_dir.iterdir():
            if d.is_dir():
                rel_path = d.relative_to(NOTES).as_posix() + "/"
                folders.append(rel_path)
        folders = sorted(folders)
        
        return {"files": files, "folders": folders, "current_folder": folder}


@app.get("/files/{name:path}")
async def read_file(name: str):
    """Read a file by its path (can include folder path)."""
    path = validate_path(name)
    if not path.exists():
        raise HTTPException(404, "File not found")
    if path.is_dir():
        raise HTTPException(400, "Path is a directory, not a file")
    return {"name": name, "content": path.read_text()}


@app.post("/files/{name:path}")
async def save_file(name: str, body: NoteBody):
    """Save a file by its path (can include folder path)."""
    validate_path(name)
    path = NOTES / name
    
    # Handle folder creation - paths ending with /.folder or just folder path
    if name.endswith("/.folder") or (not name.endswith(".md") and not "." in Path(name).name):
        folder_path = NOTES / name.replace("/.folder", "")
        folder_path.mkdir(parents=True, exist_ok=True)
        return {"saved": name, "type": "folder"}
    
    # Regular file save
    if not name.endswith(".md"):
        name += ".md"
        path = NOTES / name
    
    # Ensure parent directory exists
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(body.content)
    return {"saved": name, "type": "file"}


@app.delete("/files/{name:path}")
async def delete_file(name: str):
    """Delete a file or folder by its path."""
    path = validate_path(name)
    if not path.exists():
        raise HTTPException(404, "File or folder not found")
    
    if path.is_dir():
        # Delete folder and all contents
        shutil.rmtree(path)
        return {"deleted": name, "type": "folder"}
    else:
        path.unlink()
        return {"deleted": name, "type": "file"}


# ── folder operations ──────────────────────────────────────────────────────────

@app.post("/folders/{name:path}")
async def create_folder(name: str):
    """Create a new folder."""
    folder_path = NOTES / name
    folder_path.mkdir(parents=True, exist_ok=True)
    return {"created": name}


@app.put("/folders/{name:path}")
async def rename_folder(name: str, new_name: str = ""):
    """Rename a folder."""
    if not new_name:
        raise HTTPException(400, "new_name parameter required")
    old_path = NOTES / name
    new_path = NOTES / new_name
    if not old_path.exists():
        raise HTTPException(404, f"Folder '{name}' not found")
    old_path.rename(new_path)
    return {"renamed": name, "to": new_name}


@app.post("/reveal/{name:path}")
async def reveal_in_file_manager(name: str):
    """Open the file manager with the given path selected (Linux xdg-open)."""
    import subprocess
    target = NOTES / name
    if not target.exists():
        raise HTTPException(404, f"'{name}' not found")
    # Open the parent directory so the item is visible
    parent = target.parent if target.is_file() else target
    subprocess.Popen(["xdg-open", str(parent)])
    return {"opened": str(parent)}


# ── search ─────────────────────────────────────────────────────────────────────

@app.get("/search")
async def search_notes(q: str = ""):
    """Full-text search across all notes (recursive). Returns [{name, snippet}]."""
    if not q:
        return {"results": []}
    q_lower = q.lower()
    results = []
    # Search recursively for all .md files
    for path in sorted(NOTES.rglob("*.md")):
        content = path.read_text(errors="replace")
        if q_lower in content.lower():
            # find first matching line for snippet
            for line in content.splitlines():
                if q_lower in line.lower():
                    snippet = line.strip()[:120]
                    break
            else:
                snippet = content[:120]
            # Return path relative to NOTES
            rel_path = path.relative_to(NOTES).as_posix()
            results.append({"name": rel_path, "snippet": snippet})
    return {"results": results}


# ── download / export ──────────────────────────────────────────────────────────

@app.get("/download/md/{name}")
async def download_md(name: str):
    path = NOTES / name
    if not path.exists():
        raise HTTPException(404)
    return FileResponse(path, media_type="text/markdown",
                        headers={"Content-Disposition": f'attachment; filename="{name}"'})


@app.post("/export/html/{name}")
async def export_html(name: str, body: HtmlBody):
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{name}</title>
<style>
  body {{ font-family: Georgia, serif; max-width: 800px; margin: 2rem auto; padding: 0 1rem; line-height: 1.6; }}
  pre {{ background: #f4f4f4; padding: 1rem; border-radius: 4px; overflow-x: auto; }}
  code {{ background: #f4f4f4; padding: .2em .4em; border-radius: 3px; }}
  blockquote {{ border-left: 4px solid #ccc; margin: 0; padding-left: 1rem; color: #555; }}
  img {{ max-width: 100%; }}
</style>
</head>
<body>
{body.html}
</body>
</html>"""
    stem = name.removesuffix(".md")
    return Response(
        content=html,
        media_type="text/html",
        headers={"Content-Disposition": f'attachment; filename="{stem}.html"'},
    )


@app.post("/export/pdf/{name}")
async def export_pdf(name: str, body: HtmlBody):
    WP_HTML = _get_weasyprint()
    if not WP_HTML:
        raise HTTPException(501, "WeasyPrint not installed")
    stem = name.removesuffix(".md")
    pdf_bytes = WP_HTML(string=body.html).write_pdf()
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{stem}.pdf"'},
    )


# ── image upload ───────────────────────────────────────────────────────────────

IMAGES = BASE / "static" / "images"
IMAGES.mkdir(exist_ok=True)

@app.post("/upload/image")
async def upload_image(file: UploadFile = File(...)):
    filename = Path(file.filename or "image.png").name
    # sanitise filename
    safe = "".join(c for c in filename if c.isalnum() or c in "._-")
    if not safe:
        safe = "image.png"
    dest = IMAGES / safe
    # avoid collisions
    stem, ext = os.path.splitext(safe)
    counter = 1
    while dest.exists():
        dest = IMAGES / f"{stem}_{counter}{ext}"
        counter += 1
    dest.write_bytes(await file.read())
    return {"url": f"/static/images/{dest.name}"}


# ── import ─────────────────────────────────────────────────────────────────────

@app.post("/import")
async def import_file(file: UploadFile = File(...)):
    filename = file.filename or ""
    ext = Path(filename).suffix.lower()
    data = await file.read()

    if ext in (".md", ".txt"):
        return {"content": data.decode("utf-8", errors="replace")}

    if ext == ".pdf":
        fitz = _get_fitz()
        if not fitz:
            raise HTTPException(501, "PyMuPDF not installed — run: pip install pymupdf")
        content = _pdf_to_markdown(fitz, data)
        return {"content": content}

    if ext == ".docx":
        DocxDocument = _get_docx()
        if not DocxDocument:
            raise HTTPException(501, "python-docx not installed — run: pip install python-docx")
        doc = DocxDocument(io.BytesIO(data))
        lines = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            if style.startswith("Heading 1"):
                lines.append(f"# {para.text}")
            elif style.startswith("Heading 2"):
                lines.append(f"## {para.text}")
            elif style.startswith("Heading 3"):
                lines.append(f"### {para.text}")
            else:
                lines.append(para.text)
        return {"content": "\n\n".join(lines)}

    raise HTTPException(415, f"Unsupported file type: {ext}")


# ── github publish ─────────────────────────────────────────────────────────────

@app.post("/publish/{name}")
async def publish(name: str, body: PublishBody):
    git = _get_git()
    if not git:
        raise HTTPException(501, "gitpython not installed")
    cfg = load_config()
    gh = cfg.get("github", {})
    repo_url = gh.get("repo_url", "")
    token = gh.get("token", "")
    branch = gh.get("branch", "main")

    if not repo_url:
        raise HTTPException(400, "GitHub repo_url not configured in config.json")

    if not name.endswith(".md"):
        name += ".md"
    note_path = NOTES / name
    note_path.write_text(body.content)

    authed_url = repo_url.replace("https://", f"https://{token}@") if (token and "https://" in repo_url) else repo_url

    repo_dir = BASE / ".git_repo_cache"
    try:
        if repo_dir.exists():
            repo = git.Repo(repo_dir)
            repo.remotes.origin.set_url(authed_url)
            repo.remotes.origin.pull()
        else:
            repo = git.Repo.clone_from(authed_url, repo_dir, branch=branch)

        shutil.copy(note_path, repo_dir / name)
        repo.index.add([name])
        repo.index.commit(body.message or f"Update {name}")
        repo.remotes.origin.set_url(authed_url)
        repo.remotes.origin.push(branch)
    except Exception as e:
        raise HTTPException(500, f"Git error: {e}")

    results = {"published": name, "branch": branch}

    # ── optional Google Drive sync ────────────────────────────────────────────
    gd_cfg = cfg.get("gdrive", {})
    if gd_cfg.get("enabled") and GDRIVE_TOKEN_PATH.exists():
        gd = _get_gdrive()
        if not gd:
            results["gdrive"] = "google SDK not installed"
        else:
            try:
                creds = gd["Credentials"].from_authorized_user_file(str(GDRIVE_TOKEN_PATH), GDRIVE_SCOPES)
                service = gd["build"]("drive", "v3", credentials=creds)
                file_metadata = {"name": name}
                media = gd["MediaIoBaseUpload"](io.BytesIO(note_path.read_bytes()), mimetype="text/markdown")
                # check if file already exists
                existing = service.files().list(
                    q=f"name='{name}' and trashed=false",
                    fields="files(id)"
                ).execute().get("files", [])
                if existing:
                    service.files().update(fileId=existing[0]["id"], media_body=media).execute()
                else:
                    service.files().create(body=file_metadata, media_body=media, fields="id").execute()
                results["gdrive"] = "synced"
            except Exception as e:
                results["gdrive"] = f"error: {e}"

    return results


# ── google drive oauth ─────────────────────────────────────────────────────────

@app.get("/gdrive/auth")
async def gdrive_auth(json: bool = False):
    """Start Google OAuth2 flow."""
    gd = _get_gdrive()
    if not gd:
        raise HTTPException(501, "google SDK not installed")
    flow = gd["Flow"].from_client_config(
        GDRIVE_CLIENT_CONFIG,
        scopes=GDRIVE_SCOPES,
        redirect_uri="http://localhost:8000/gdrive/callback",
        autogenerate_code_verifier=False,
    )
    auth_url, _ = flow.authorization_url(prompt="consent", access_type="offline")
    if json:
        return {"auth_url": auth_url}
    return RedirectResponse(auth_url)


@app.get("/gdrive/callback")
async def gdrive_callback(code: str = None, state: str = "", error: str = None):
    if error:
        return HTMLResponse(f"<html><body><h2>Google Drive auth failed</h2><p>{error}</p></body></html>")
    if not code:
        return HTMLResponse("<html><body><h2>Google Drive auth failed</h2><p>Missing code</p></body></html>")
    gd = _get_gdrive()
    if not gd:
        raise HTTPException(501, "google SDK not installed")
    try:
        flow = gd["Flow"].from_client_config(
            GDRIVE_CLIENT_CONFIG,
            scopes=GDRIVE_SCOPES,
            redirect_uri="http://localhost:8000/gdrive/callback",
            autogenerate_code_verifier=False,
        )
        flow.fetch_token(code=code)
        creds = flow.credentials
        GDRIVE_TOKEN_PATH.write_text(creds.to_json())
        # enable gdrive in config
        cfg = load_config()
        cfg.setdefault("gdrive", {})["enabled"] = True
        save_config(cfg)
    except Exception as e:
        logger.error(f"Google Drive OAuth callback error: {e}")
        return HTMLResponse(f"<html><body><h2>Google Drive auth failed</h2><p>{e}</p></body></html>")
    return HTMLResponse("<html><body><h2>Google Drive connected!</h2><p>You can close this tab.</p><script>window.close()</script></body></html>")


@app.get("/gdrive/status")
async def gdrive_status():
    connected = GDRIVE_TOKEN_PATH.exists()
    return {"connected": connected}


@app.get("/gdrive/files")
async def gdrive_list():
    gd = _get_gdrive()
    if not gd:
        raise HTTPException(501, "google SDK not installed")
    if not GDRIVE_TOKEN_PATH.exists():
        raise HTTPException(401, "Not authenticated with Google Drive")
    try:
        creds = gd["Credentials"].from_authorized_user_file(str(GDRIVE_TOKEN_PATH), GDRIVE_SCOPES)
        service = gd["build"]("drive", "v3", credentials=creds)
        results = service.files().list(
            q="name contains '.md' and trashed=false",
            fields="files(id,name)"
        ).execute()
        return {"files": [f["name"] for f in results.get("files", [])]}
    except Exception as e:
        raise HTTPException(500, str(e))


@app.get("/gdrive/open/{name}")
async def gdrive_open(name: str):
    gd = _get_gdrive()
    if not gd:
        raise HTTPException(501, "google SDK not installed")
    if not GDRIVE_TOKEN_PATH.exists():
        raise HTTPException(401, "Not authenticated with Google Drive")
    try:
        creds = gd["Credentials"].from_authorized_user_file(str(GDRIVE_TOKEN_PATH), GDRIVE_SCOPES)
        service = gd["build"]("drive", "v3", credentials=creds)
        results = service.files().list(
            q=f"name='{name}' and trashed=false",
            fields="files(id,name)"
        ).execute().get("files", [])
        if not results:
            raise HTTPException(404, f"{name} not found on Google Drive")
        file_id = results[0]["id"]
        buf = io.BytesIO()
        request = service.files().get_media(fileId=file_id)
        downloader = gd["MediaIoBaseDownload"](buf, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        return {"name": name, "content": buf.getvalue().decode("utf-8", errors="replace")}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))


# ── GitHub OAuth ───────────────────────────────────────────────────────────────

@app.get("/auth/github/client-id")
async def github_client_id():
    """Return the OAuth client ID so the frontend can initiate the flow."""
    return {"client_id": GITHUB_CLIENT_ID}


@app.get("/github/status")
async def github_status():
    cfg = load_config()
    gh = cfg.get("github", {})
    return {
        "connected": GITHUB_TOKEN_PATH.exists(),
        "repo_url": gh.get("repo_url", ""),
        "branch": gh.get("branch", "main"),
    }


@app.get("/github/repos")
async def github_repos():
    """List the authenticated user's repositories."""
    cfg = load_config()
    token = cfg.get("github", {}).get("token", "")
    if not token:
        raise HTTPException(401, "Not signed in to GitHub")
    httpx = _get_httpx()
    repos = []
    page = 1
    async with httpx.AsyncClient() as client:
        while True:
            resp = await client.get(
                "https://api.github.com/user/repos",
                params={"per_page": 100, "sort": "updated", "page": page},
                headers={"Authorization": f"token {token}", "Accept": "application/vnd.github+json"},
            )
            if resp.status_code != 200:
                raise HTTPException(resp.status_code, "Failed to fetch repos")
            batch = resp.json()
            if not batch:
                break
            for r in batch:
                repos.append({
                    "name": r["name"],
                    "full_name": r["full_name"],
                    "clone_url": r["clone_url"],
                    "default_branch": r["default_branch"],
                    "private": r["private"],
                })
            page += 1
    return {"repos": repos}


class SelectRepoRequest(BaseModel):
    repo_url: str
    branch: str = "main"


@app.post("/github/select-repo")
async def github_select_repo(request: SelectRepoRequest):
    """Select a repo, save to config, and clone it."""
    cfg = load_config()
    token = cfg.get("github", {}).get("token", "")
    if not token:
        raise HTTPException(401, "Not signed in to GitHub")

    cfg.setdefault("github", {})
    cfg["github"]["repo_url"] = request.repo_url
    cfg["github"]["branch"] = request.branch
    save_config(cfg)

    # Clone or update the repo cache
    git = _get_git()
    if not git:
        raise HTTPException(501, "gitpython not installed")

    authed_url = request.repo_url.replace("https://", f"https://{token}@")
    cache = BASE / ".git_repo_cache"
    try:
        if cache.exists():
            repo = git.Repo(cache)
            repo.remotes.origin.set_url(authed_url)
            repo.remotes.origin.pull()
        else:
            git.Repo.clone_from(authed_url, cache, branch=request.branch)
    except Exception as e:
        raise HTTPException(500, f"Clone failed: {e}")

    return {"status": "ok", "repo_url": request.repo_url, "branch": request.branch}


# ── Publish to Cloud ───────────────────────────────────────────────────────────

@app.post("/publish")
async def publish_all():
    """Publish all notes to connected cloud service."""
    # Check which service is connected
    github_connected = GITHUB_TOKEN_PATH.exists()
    gdrive_connected = GDRIVE_TOKEN_PATH.exists()
    
    if not any([github_connected, gdrive_connected]):
        raise HTTPException(400, "No cloud service connected. Sign in to GitHub or Google Drive first.")
    
    # Collect all files
    files = {}
    for p in NOTES.rglob("*.md"):
        rel_path = p.relative_to(NOTES).as_posix()
        files[rel_path] = p.read_bytes()
    
    if not files:
        raise HTTPException(400, "No notes to publish")
    
    results = {}
    
    # Publish to GitHub
    git = _get_git()
    if github_connected and git:
        try:
            token_data = json.loads(GITHUB_TOKEN_PATH.read_text())
            token = token_data["access_token"]
            cfg = load_config().get("github", {})
            repo_url = cfg.get("repo_url")
            branch = cfg.get("branch", "main")
            
            if not repo_url:
                results["github"] = "Error: repo_url not configured"
            else:
                cache = BASE / ".git_cache"
                if cache.exists():
                    shutil.rmtree(cache)
                auth_url = repo_url.replace("https://", f"https://{token}@")
                repo = git.Repo.clone_from(auth_url, cache, branch=branch)
                
                for rel_path, content in files.items():
                    dest = cache / rel_path
                    dest.parent.mkdir(parents=True, exist_ok=True)
                    dest.write_bytes(content)
                
                repo.git.add(A=True)
                if repo.is_dirty():
                    repo.index.commit(f"Publish {len(files)} notes")
                    repo.remote().push()
                    results["github"] = f"✅ Pushed {len(files)} files"
                else:
                    results["github"] = "✅ Up to date"
        except Exception as e:
            results["github"] = f"Error: {e}"
    
    # Publish to Google Drive
    gd = _get_gdrive()
    if gdrive_connected and gd:
        try:
            creds = gd["Credentials"].from_authorized_user_file(str(GDRIVE_TOKEN_PATH), GDRIVE_SCOPES)
            service = gd["build"]("drive", "v3", credentials=creds)

            # Find or create root folder
            results_list = service.files().list(q="name='Lectura' and mimeType='application/vnd.google-apps.folder' and trashed=false", fields="files(id)").execute()
            if results_list.get("files"):
                root_id = results_list["files"][0]["id"]
            else:
                root_meta = {"name": "Lectura", "mimeType": "application/vnd.google-apps.folder"}
                root_folder = service.files().create(body=root_meta, fields="id").execute()
                root_id = root_folder["id"]
            
            # Helper to get or create folder
            folder_cache = {".": root_id}
            def get_folder_id(path):
                if path in folder_cache:
                    return folder_cache[path]
                parent_path = "/".join(path.split("/")[:-1]) or "."
                parent_id = get_folder_id(parent_path)
                folder_name = path.split("/")[-1]
                q = f"name='{folder_name}' and '{parent_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"
                res = service.files().list(q=q, fields="files(id)").execute()
                if res.get("files"):
                    folder_id = res["files"][0]["id"]
                else:
                    meta = {"name": folder_name, "mimeType": "application/vnd.google-apps.folder", "parents": [parent_id]}
                    folder = service.files().create(body=meta, fields="id").execute()
                    folder_id = folder["id"]
                folder_cache[path] = folder_id
                return folder_id
            
            # Upload files preserving folder structure
            for rel_path, content in files.items():
                parts = rel_path.split("/")
                if len(parts) > 1:
                    folder_path = "/".join(parts[:-1])
                    parent_id = get_folder_id(folder_path)
                else:
                    parent_id = root_id
                
                file_name = parts[-1]
                media = gd["MediaIoBaseUpload"](io.BytesIO(content), mimetype="text/markdown")
                file_meta = {"name": file_name, "parents": [parent_id]}
                service.files().create(body=file_meta, media_body=media).execute()
            
            results["gdrive"] = f"✅ Uploaded {len(files)} files to Lectura/"
        except Exception as e:
            results["gdrive"] = f"Error: {e}"
    
    return {"published": len(files), "results": results}


# ── config ─────────────────────────────────────────────────────────────────────

@app.get("/config")
async def get_config():
    cfg = load_config()
    if "github" in cfg:
        cfg["github"]["token"] = "***" if cfg["github"].get("token") else ""
    gdrive_connected = GDRIVE_TOKEN_PATH.exists()
    cfg.setdefault("gdrive", {})["connected"] = gdrive_connected
    return cfg


@app.post("/config")
async def post_config(body: ConfigBody):
    existing = load_config()
    if body.config.get("github", {}).get("token") == "***":
        body.config["github"]["token"] = existing.get("github", {}).get("token", "")
    # preserve gdrive connected state
    body.config.setdefault("gdrive", {}).pop("connected", None)
    save_config(body.config)
    return {"saved": True}


# ═══════════════════════════════════════════════════════════════════════════════
# GITHUB OAUTH ENDPOINTS
# ═══════════════════════════════════════════════════════════════════════════════

# GitHub OAuth token exchange

class GitHubTokenRequest(BaseModel):
    code: str

@app.post("/auth/github/token")
async def github_token_exchange(request: GitHubTokenRequest):
    """Exchange GitHub OAuth code for access token"""
    try:
        httpx = _get_httpx()
        async with httpx.AsyncClient() as client:
            # Exchange code for access token
            token_response = await client.post(
                'https://github.com/login/oauth/access_token',
                data={
                    'client_id': GITHUB_CLIENT_ID,
                    'client_secret': GITHUB_CLIENT_SECRET,
                    'code': request.code,
                },
                headers={'Accept': 'application/json'}
            )
            
            if token_response.status_code != 200:
                raise HTTPException(status_code=400, detail="Token exchange failed")
            
            token_data = token_response.json()
            
            if 'error' in token_data:
                raise HTTPException(status_code=400, detail=token_data.get('error_description', 'OAuth error'))
            
            access_token = token_data.get('access_token')
            # Save token to config so git operations can use it
            cfg = load_config()
            cfg.setdefault("github", {})["token"] = access_token
            save_config(cfg)
            GITHUB_TOKEN_PATH.write_text(json.dumps(token_data))
            return {"access_token": access_token}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"GitHub OAuth error: {e}")
        raise HTTPException(status_code=500, detail="OAuth request failed")

@app.get("/auth/github/callback")
async def github_callback(code: str = None, state: str = None, error: str = None):
    """Handle GitHub OAuth callback — works both as popup and system browser redirect."""
    if error:
        return HTMLResponse(f"<html><body><h2>GitHub auth failed</h2><p>{error}</p></body></html>")

    if not code or not state:
        return HTMLResponse("<html><body><h2>GitHub auth failed</h2><p>Missing code or state</p></body></html>")

    # Exchange token server-side so it works from the system browser too
    try:
        httpx = _get_httpx()
        async with httpx.AsyncClient() as client:
            token_response = await client.post(
                'https://github.com/login/oauth/access_token',
                data={
                    'client_id': GITHUB_CLIENT_ID,
                    'client_secret': GITHUB_CLIENT_SECRET,
                    'code': code,
                },
                headers={'Accept': 'application/json'}
            )
            token_data = token_response.json()
            if 'error' in token_data:
                return HTMLResponse(f"<html><body><h2>GitHub auth failed</h2><p>{token_data.get('error_description', 'OAuth error')}</p></body></html>")
            access_token = token_data.get('access_token')
            cfg = load_config()
            cfg.setdefault("github", {})["token"] = access_token
            save_config(cfg)
            GITHUB_TOKEN_PATH.write_text(json.dumps(token_data))
    except Exception as e:
        logger.error(f"GitHub OAuth callback error: {e}")
        return HTMLResponse(f"<html><body><h2>GitHub auth failed</h2><p>{e}</p></body></html>")

    return HTMLResponse("""<html><body>
        <h2>GitHub connected!</h2><p>You can close this tab.</p>
        <script>
            if (window.opener) {
                window.opener.postMessage({ type: 'github-auth-success' }, '*');
            }
            window.close();
        </script>
    </body></html>""")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="warning")
